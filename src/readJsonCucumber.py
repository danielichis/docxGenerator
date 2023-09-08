import os
import json
import shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image
from pathlib import Path
from datetime import datetime
import sys
class GenWords:
    def __init__(self,mode="out"):
        self.auto_path=None
        self.testMode=None
        self.appPath=None
        self.indexJson=None
        self.clearMode=mode
        self.OnedrivePath=None
        self.get_app_path()
        self.readConfig()
        self.serenityFolderPath=None
        self.detailLevel=None
        self.targetPath=None
        self.get_serenityFolderPath()
        self.JsonPath=None
        self.search_json()
        self.JsonData=None
        self.read_json()
        self.filteredJson_1=self.filtered_json_level_1()
        self.TemplateDoc=os.path.join(self.folderProyect,"Archivos Plantilla/EPC_EmptyTemplate.docx")
    def get_serenityFolderPath(self):
        if self.testMode=="TRUE":
            self.serenityFolderPath=os.path.join(self.auto_path, "target","site","serenity")
            print(self.serenityFolderPath)
        elif self.testMode=="FALSE":
            print("ruta de la carpeta serenity:")
            print(self.appPath)
            self.targetPath=os.path.dirname(os.path.dirname(self.appPath))
            self.serenityFolderPath=os.path.join(self.targetPath,"site","serenity")
            print(self.serenityFolderPath)
    def get_app_path(self):
        if getattr(sys, 'frozen', False):
            application_path = os.path.dirname(sys.executable)
        elif __file__:
            application_path = os.path.dirname(__file__)
        self.appPath =Path(application_path)
        self.folderProyect=self.appPath.parent.absolute()
        
        #return Path(application_path)


    def readConfig(self):
        with open(os.path.join(self.folderProyect,"generarWordsConfig.txt")) as f:
            lines=f.readlines()
          
        self.OnedrivePath=lines[0].split("=")[1].strip()
        self.detailLevel=lines[1].split("=")[1].strip()
        self.auto_path=lines[2].split("=")[1].strip()
        self.testMode=lines[3].split("=")[1].strip()
        self.indexJson=lines[4].split("=")[1].strip()
        self.clearMode=lines[5].split("=")[1].strip()

    def search_json(self):
        listFiles=[]
        for file in os.listdir(self.serenityFolderPath):
            if file.endswith(".json"):
                filePath=os.path.join(self.serenityFolderPath, file)
                #get the date of the file
                dateFile=os.path.getmtime(filePath)
                objectDate=datetime.fromtimestamp(dateFile)
                sizeFile=os.path.getsize(filePath)
                listFiles.append([filePath,dateFile,sizeFile])
        
        #get the last file
        listFiles.sort(key=lambda x: x[1],reverse=True)
        if len(listFiles)==0:
            print("No se encontraron archivos json")
            return
        self.JsonPath=listFiles[int(self.indexJson)][0]
    def delete_json(self):
        for file in os.listdir(self.serenityFolderPath):
            if file.endswith(".json"):
                filePath=os.path.join(self.serenityFolderPath, file)
                os.remove(filePath)
    def read_json(self):
        #read json file
        with open(self.JsonPath) as json_file:
            data = json.load(json_file)
            self.JsonData=data

    def filtered_json_level_1(self):
        casos=self.JsonData['testSteps']
        dictData={}
        dictData['LevelDetail']=self.detailLevel
        dictData['TestCases']={}
        
        for i,caso in enumerate(casos):
            dictData['TestCases'][i+1]={}
            #if caso['result']=='SUCCESS':
            testSteps=caso['children']
            dictData['TestCases'][i+1]['TestSteps']={}
            dictData['TestCases'][i+1]['testCaseDescription']=""
            dictData['TestCases'][i+1]['startTime']=caso['startTime']

            for j,testStep in enumerate(testSteps):
                dictData['TestCases'][i+1]['TestSteps'][j+1]={}
                fileImageName=f'CP-{str(i+1).zfill(3)} EVIDENCIA NRO {str(j+1).zfill(3)}'
                if testStep.keys().__contains__('screenshots'):
                    sc=testStep['screenshots'][-1]['screenshotName']
                    sc=os.path.join(self.serenityFolderPath,sc)
                    dictData['TestCases'][i+1]['TestSteps'][j+1]['name']=fileImageName
                    dictData['TestCases'][i+1]['TestSteps'][j+1]['descripcion']=testStep['description']
                    #copiar archivo a otra carpeta
                    fileImageNamePath=os.path.join(self.folderProyect,f'Archivos Input/{fileImageName}.png')
                    shutil.copy(sc,fileImageNamePath)

                    if j==0:
                        sep=""
                    else:
                        sep=" "
                    dictData['TestCases'][i+1]['testCaseDescription']=dictData['TestCases'][i+1]['testCaseDescription']+sep+testStep['description']
                else:
                    print("no se encontro imagen")
        with open(os.path.join(self.folderProyect,"data.json"), 'w') as outfile:
            json.dump(dictData, outfile, indent=4)

    def testCase_word(self):
        if self.testCaseDescription=="":
            print("no se encontro descripcion")
            return
        doc=Document(self.TemplateDoc)
        tables = doc.tables
        p = tables[0].rows[0].cells[0].add_paragraph(self.testCaseNumber)
        p = tables[0].rows[0].cells[1].add_paragraph("DANIEL CHACON")
        p = tables[0].rows[0].cells[2].add_paragraph(self.testCaseDate)
        p = tables[0].rows[1].cells[0].add_paragraph(self.testCaseDescription)
        for j,testStep in enumerate(self.testSteps):
            imagePath=f'Archivos Input/{self.testSteps[str(j+1)]["name"]}.png'
            imageFullPath=os.path.join(self.folderProyect,imagePath)
            pic=Image.open(imageFullPath)
            w=pic.size[0]
            h=pic.size[1]
            if w<300:
                w= Inches(4.5)
                h=Inches(h/96*2)
            else:
                w= Inches(7)
                h= Inches(3)
            if j==0:
                sep=""
            else:
                sep="\n\n"
            p = tables[0].rows[2].cells[0].add_paragraph(sep)
            r = p.add_run()
            r.add_text(self.testSteps[str(j+1)]['descripcion'])
            r.add_picture(imageFullPath,width=w)
        for table in doc.tables:
            for cell in table._cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name=f"{self.testCaseNumber}.docx"
        wordFilePath=os.path.join(self.OnedrivePath,name)
        doc.save(wordFilePath)
    def getDateCase(self):
        strDate=self.testCaseDate
        strDate=strDate[:19]
        objetDate=datetime.strptime(strDate, '%Y-%m-%dT%H:%M:%S')
        newFormat=objetDate.strftime("%d/%m/%Y %H:%M:%S")
        self.testCaseDate=newFormat
    def generateWords(self):
        #leer json file
        print("GENERANDO ARCHIVOS WORD")
        if self.clearMode=="TRUE":
            self.delete_json()
        with open(os.path.join(self.folderProyect,"data.json")) as json_file:
            data = json.load(json_file)
        cases=data['TestCases']
        for i,case in enumerate(cases):
            self.testSteps=cases[str(i+1)]['TestSteps']
            self.testCaseDescription=cases[str(i+1)]['testCaseDescription']
            self.testCaseNumber=f"EPC_{str(i+1).zfill(3)}"
            self.testCaseDate=cases[str(i+1)]['startTime']
            self.getDateCase()
            self.testCase_word()
gw=GenWords()
gw.generateWords()


