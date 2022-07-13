import os
import time
import numpy as np
import pyperclip

from io import StringIO
from datetime import datetime

from docx import Document
from docx.shared import Inches
from pywinauto import keyboard, Application
from Config.Parameters import *
from Config import RobotParameters
from Interactions.Apps.Desktop.SNA32.BaseSNA import copy_all_value_in_screen_command
from Interactions.Apps.Desktop.SNA32.Connex.BaseConnex import *
from Interactions.Os.Windows.KeyboardActions import press_enter, press_tab
from Utils.DateConverter import convert_date


def open_fti_application():
    status = navigate_to_option(int(FTI_TABS_TO_NAVIGATE), 5, FTI_INTEFACE_VAUE_TO_CHECK)
    if not status:
        raise
    else:
        return status


def get_values():
    result = {}
    pan = RobotParameters.values()["PAN"]
    if len(pan) != 16:
        raise
    dates_formatted = {}
    date_range_start = RobotParameters.values()["StartDate"]
    date_range_end = RobotParameters.values()["EndDate"]
    only_first = RobotParameters.values()["GetFirst"]
    details = RobotParameters.values()["Details"]
    name_directory = RobotParameters.user()["ReportDirectory"] + RobotParameters.values()["NameDirectory"]
    os.mkdir(name_directory)
    require_first = False
    if (date_range_start != "None" or date_range_start != "") and (
            date_range_end != "None" or date_range_end != ""):
        values = {0: date_range_start, 1: date_range_end}
        dates_formatted = convert_date(values)
    if only_first:
        require_first = True
    expected_value = RobotParameters.values()["ExpectedValue"]
    result = {
        "card_number": pan,
        "range_between": dates_formatted,
        "details": details,
        "expected_value": expected_value,
        "require_first": require_first,
        "name_directory": name_directory
    }
    return result


def consult_transaction():
    try:
        all_values = get_values()
        card_number = all_values["card_number"]
        range_between = all_values["range_between"]
        require_detail = all_values["details"]
        only_first = all_values["require_first"]
        name_directory = all_values["name_directory"]
        expected_value = all_values["expected_value"]
        if name_directory is None:
            raise
        time.sleep(3)
        keyboard.send_keys(card_number)
        if range_between is not None:
            press_tab(multiple_times=16)
            start_date = range_between[0]
            end_date = range_between[1]
            keyboard.send_keys(start_date)
            press_tab()
            keyboard.send_keys(end_date)
            print("paso fecha")
        press_enter()
        time.sleep(2)
        data = general_transaction_values(name_directory)
        document = Document()
        document.add_heading('Resumen de Ejecucion' + str(card_number), 0)
        document.add_paragraph("------------ Transaccion(es) tarjeta: " + str(card_number) + " ------------")
        print("paso documento inicial")
        if range_between is None:
            now = datetime.now()
            dt_string = now.strftime("%d/%m/%Y")
            document.add_paragraph("------ Hoy: " + dt_string + " ------")
        else:
            document.add_paragraph("------ Desde: " + range_between[2] + " Hasta: " + range_between[3] + " ------")
        print("paso documento segundo")
        app = Application(backend="uia").connect(path="C:\Program Files (x86)\SNA95\system\WIN3270.EXE")
        document.add_paragraph(str(pyperclip.paste()))
        take_screenshot(app, document)
        detail_text = ""
        list_data = data[0]
        loops = data[1]
        for x in range(loops):
            if x != 0:
                keyboard.send_keys("{F3}")
                time.sleep(2)
            time.sleep(2)
            press_tab(multiple_times=x + 2)
            keyboard.send_keys("x")
            press_enter()
            time.sleep(2)
            document.add_paragraph(" --------- Fecha: " + list_data[x]["FormattedDate"] + " Hora: " + list_data[x][
                "FormattedTime"] + " ---------")
            copy_all_value_in_screen_command()
            if list_data[x]["RC"] != expected_value:
                document.add_paragraph(
                    "--------------------- ATENCION: La transaccion no cumple con el valor esperado:" + expected_value + "  ---------------------")
                document.add_paragraph(
                    "--------------------- el valor de la transacción es:" + list_data[x][
                        "RC"] + "  ---------------------")
            document.add_paragraph(str(pyperclip.paste()))
            document.add_paragraph("--------------------- FIN ---------------------")
            take_screenshot(app, document)
            if require_detail == "Detalles Totales":
                press_enter()
                time.sleep(2)
                take_screenshot(app, document)
                press_enter()
                time.sleep(2)
                take_screenshot(app, document)
                press_enter()
                time.sleep(2)
                take_screenshot(app, document)
                press_enter()
                time.sleep(2)
                take_screenshot(app, document)
            if only_first:
                break
        final = {0: True, 1: list_data}
        document.save(name_directory + '\\final_report.docx')
        os.startfile('"' + name_directory + '\\final_report.docx"')
        return final
    except Exception as err:
        final = {0: False, 1: str(err)}
        # creating/opening a file
        f = open("logger.txt", "a")

        # writing in the file
        f.write(str(err))

        # closing the file
        f.close()
        return final


def general_transaction_values(name_directory):
    copy_all_value_in_screen_command()
    s = str(pyperclip.paste())
    with open(name_directory + '\\general_transaction_values.txt', 'w') as g:
        g.write(s)
    a_file = open(name_directory + "\\general_transaction_values.txt", "r")
    lines = a_file.readlines()
    a_file.close()
    del lines[0:12]
    del lines[32:37]
    text = ''.join(str(e) for e in lines)
    c = StringIO(text)
    data = np.loadtxt(c, dtype='str')
    length = len(data)
    result = []
    if (isinstance(data[0], str)):
        line_values = {}
        line_values["CardholderNumber/PAN"] = data[1]
        line_values["Terminal"] = data[2]
        line_values["Date"] = data[3]
        line_values["FormattedDate"] = date_hour_formatter(data[3], "date")
        line_values["Time"] = data[4]
        line_values["FormattedTime"] = date_hour_formatter(data[4], "hour")
        line_values["P-CODE"] = data[5]
        line_values["SI"] = data[6]
        line_values["RC"] = data[7]
        line_values["Amount"] = data[8]
        result.append(line_values)
        length = 1
    else:
        for value in data:
            line_values = {}
            line_values["CardholderNumber/PAN"] = value[1]
            line_values["Terminal"] = value[2]
            line_values["Date"] = value[3]
            line_values["FormattedDate"] = date_hour_formatter(value[3], "date")
            line_values["Time"] = value[4]
            line_values["FormattedTime"] = date_hour_formatter(value[4], "hour")
            line_values["P-CODE"] = value[5]
            line_values["SI"] = value[6]
            line_values["RC"] = value[7]
            line_values["Amount"] = value[8]
            result.append(line_values)
    return result, length


def expected_transaction_value_comparer(final_values, expected_value, name_directory):
    try:
        values = final_values
        log_file_value = " --------- Los siguientes movimientos no tienen el valor esperado de la transaccion --------- \n"
        counter = 0
        failed_values = ""
        for item in values:
            if item["RC"] != expected_value:
                failed_values = failed_values + " ---------------------------- \n"
                failed_values = failed_values + "\n"
                failed_values = failed_values + "PAN:" + item["CardholderNumber/PAN"] + ";"
                failed_values = failed_values + "Terminal:" + item["Terminal"] + ";"
                failed_values = failed_values + "FormattedDate:" + item["FormattedDate"] + ";"
                failed_values = failed_values + "FormattedTime:" + item["FormattedTime"] + ";"
                failed_values = failed_values + "P-CODE" + item["P-CODE"] + ";"
                failed_values = failed_values + "SI" + item["SI"] + ";"
                failed_values = failed_values + "RC" + item["RC"] + ";"
                failed_values = failed_values + "Monto" + item["Amount"] + ";"
                failed_values = failed_values + "\n"
                failed_values = failed_values + " ---------------------------- \n"
                counter = counter + 1
        if counter == 0:
            failed_values = failed_values + " NO SE ENCONTRARON DISCORDANCIAS "
        final_text = log_file_value + failed_values
        f = open("loggerwin.txt", "a")

        # writing in the file
        f.write(str(final_text))

        # closing the file
        f.close()
        with open(name_directory + '\\log_expected.txt', 'w') as g:
            g.write(final_text)
        return True
    except Exception as err:
        # creating/opening a file
        f = open("logger2.txt", "a")

        # writing in the file
        f.write(str(err))

        # closing the file
        f.close()
        return False


def date_hour_formatter(value, typo):
    final_value = None
    n = 2
    out = [(value[i:i + n]) for i in range(0, len(value), n)]

    if typo == "date":
        final_value = out[1] + "/" + out[0] + "/20" + out[2]
    elif typo == "hour":
        final_value = out[0] + ":" + out[1] + ":" + out[2]

    return final_value


def take_screenshot(app, document):
    hwin = app.top_window()
    hwin.set_focus()
    img = hwin.capture_as_image()
    img.save("captura.jpg")
    document.add_picture('captura.jpg', width=Inches(7), height=Inches(5))
