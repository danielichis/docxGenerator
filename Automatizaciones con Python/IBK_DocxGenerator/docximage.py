from pydoc import Doc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
import os
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from os.path import isfile, join
from PIL import Image
from tkinter import messagebox
import re

def list_imagenes(pics_folder):
    obj_images=[]
    global files_list
    files_list = [f for f in os.listdir(pics_folder) if isfile(join(pics_folder, f))]
    for file in files_list:
        filename = os.fsdecode(file)
        try:
            cpId=re.findall(r"(CP\d{3})-\d{2}-.*\.[p j][n p]g",filename)[0]
            picId=re.findall(r"CP\d{3}-(\d{2})-.*\.[p j][n p]g",filename)[0]
            all_descrp=re.findall(r"CP\d{3}-\d{2}-(.*)\.[p j][n p]g",filename)[0].replace("#","\n")
            images={
                    "CP ID":cpId,
                    "pic ID":picId,
                    "PARRRAFO":all_descrp,
                    "Entire Name":file,
                }
            obj_images.append(images)
        except:
            print(f"No se pudo agregar la imagen {filename}")
            pass
    return obj_images
def add_imagenes(AppFolder,nameDoc,modeJoin):
    mypath="Archivos Input"
    folderEpc=AppFolder+"/"+"Archivos Ouput"
    pics_folder=AppFolder+"/"+mypath

    doc = Document(nameDoc)
    tables = doc.tables
    cp_pictures=list_imagenes(pics_folder)
    image_count=0
    for i,table in enumerate(tables):
        descp_word=table.rows[1].cells[1].text
        cp_word=re.findall(r"CP\d{3}",descp_word)[0]
        for cp in cp_pictures:
            rut=cp["Entire Name"]
            #print(cp_word,cp["CP ID"])
            if cp_word==cp["CP ID"]:
                #print(f"Ingresando la imagen {cp} en el cp word {cp_word} ")
                p = table.rows[2].cells[0].add_paragraph(cp["PARRRAFO"])
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run()
                w= Image.open(f"{pics_folder}/{rut}").size[0]
                h=Image.open(f"{pics_folder}/{rut}").size[1]
                rp=w/h

                print(f"tamaño original: {w},{h}")
                if w<300:
                    w= Inches(4.5)
                    h=Inches(h/96*2)
                else:
                    w= Inches(7)
                    h= Inches(3)
                
                print(f"tamaño final: {w},{h}")
                r.add_picture(f"{pics_folder}/{rut}",width=w)#, height=h)
                image_count=image_count+1
                
    print("Finished")
    nameEndFile=nameDoc
    try:
        doc.save(nameEndFile)
        if len(files_list)==image_count or modeJoin==False:
            print (modeJoin)
            #messagebox.showinfo("PROCEDIMIENTO EXITOSO", "SE AGREGARON LAS IMAGENES :) ")
        else:
            messagebox.showwarning("Advertencia",f"ARCHIVO GENERADO \nSE AGREGARON {image_count} IMAGENES PERO HAY {len(files_list)} en carpeta {str(modeJoin)}")
    except:
        messagebox.showerror("ERROR", f"CIERRE EL ARCHIVO {nameEndFile} y vuelva a ejecutar")

#print(len(list_imagenes("Archivos Input")))