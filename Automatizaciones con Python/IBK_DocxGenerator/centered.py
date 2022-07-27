from pydoc import Doc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
import os

def aligmentTable(nameDoc):
    doc = Document(nameDoc)
    tables = doc.tables
    for i,table in enumerate(tables):
        prs =table.rows[2].cells[0].paragraphs
        for p in prs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            print(p.text)
    doc.save(nameDoc)
